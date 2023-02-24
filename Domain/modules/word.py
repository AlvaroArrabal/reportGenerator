from docx import Document
from docx.shared import Cm
import time
# 'KPI':['OK','NOK']

dic = { "2G CDR CS (%)":['Degradación puntual','Degradaciones repetitivas en '],
        "2G CSSR CS (%)":['',''],
        "2G CSSR PS (%)":['',''],
        "2G Iniciated calls":['Hay llamadas iniciadas','No hay llamadas iniciadas en '],
        "2G DL Data traffic (KB)":['',''],
        "2G UL Data traffic (KB)":['',''],
        "2G ICMBand () ":['Se cumple el target en horas valle','Valores elevados en horas valle en'],
        "2G Cell Availability (%)":['Valores correctos de availability','Problemas de availability en '],
        "2G Speech disconnections":['',''],
        "3G CDR CS (%)":['Degradación puntual','Degradaciones repetitivas en '],
        "3G CSSR CS (%)":['',''],
        "3G CSSR PS (%)":['',''],
        "3G Iniciated calls":['Hay llamadas iniciadas','No hay llamadas iniciadas en '],
        "3G DL Data traffic (KB)":['',''],
        "3G UL Data traffic (KB)":['',''],
        "3G RTWP (dBm)":['Se cumple el target en horas valle','Valores elevados en horas valle en '],
        "3G Cell Availability (%)":['Valores correctos de availability','Problemas de availability en '],
        "3G Calls ending in 2G (%)":['Valores por debajo del umbral del 10%','Valores elevados de 3G calls ending in 2G en '],
        "TH DL (2G3G4G)":['',''],
        "TH UL (2G3G4G)":['',''],
        "4G CDR (VoLTE) (%)":['Degradación puntual','Degradaciones repetitivas en '],
        "4G_DCR_DATA":['Degradación puntual','Degradaciones repetitivas en '],
        "4G CSSR (VoLTE) (%)":['',''],
        "4G CSSR PS (%)":['',''],
        "4G Iniciated calls (VoLTE)":['Hay llamadas iniciadas','No hay llamadas iniciadas en '],
        "4G DL Data traffic (MB)":['',''],
        "4G UL Data traffic (MB)":['',''],
        "4G Interference PUSCH (dBm)":['Se cumple el target en horas valle','Valores elevados en horas valle en '],
        "4G Cell Availability (%)":['Valores correctos de availability','Problemas de availability en '],
        "4G MIMO (Rank2) (%)":['Valores correctos de MIMO Rank4',''],
        "4G MIMO (Rank4) (%)":['Valores correctos de MIMO Rank4','No hay valores de MIMO Rank4 en '],
        "4G CSFB E2W":['',''],
        "4G CA in PCELL":['',''],
        "4G CA in SCELL":['',''],
        "4G IntraLTE HOSR (including preparation) ()":['',''],
        "4G SRVCC HO Att":['Hay intentos de SRVCC en','Sin intentos de SRVCC en ','Teniendo en cuenta la cantidad de llamadas iniciadas, se considera que el comportamiento es el esperado'],
        "Tput DL 4G >2Mbps":['',''],
        "Tput UL 4G >500kbps":['Valores entorno al target objetivo de TH (0.5 Mbps)','Valores bajos de TH en '],
        '4G_DCR_CS (VoLTE)':['',''],
        '4G CSSR CS (VoLTE)':['',''],
        '4G_CSSR_PS_Success_Rate':['',''],
        '5G_CSSR_PS_Success_Rate':['',''],
        '4G VoLTE Iniciated calls':['',''],       
        '4G_Downlink_Traffic_Volume_MB':['',''],
        '4G_Uplink_Traffic_Volume_MB':['',''],
        '5G_Downlink_Traffic_Volume_MB':['',''],
        '5G_Uplink_Traffic_Volume_MB':['',''],
        'Interference 4G PUSCH UL (RSSI UL 4G) ':['Se cumple el target en horas valle','Valores elevados en horas valle en '],
        'Interference 4G PUSCH UL (RSSI UL 4G)':['Se cumple el target en horas valle','Valores elevados en horas valle en '],
        'Interference 5G UL (RSSI UL 5G) *':['Se cumple el target en horas valle','Valores elevados en horas valle en '],
        '4G_Availability_Cell_Rate_Hourly':['',''],
        '5G_Availability_Cell_Rate_Hourly':['',''],
        '4G_% MIMO':['',''],
        'CSFB attempts (L.CSFB.E2W + L.CSFB.E2G)':['',''],
        'CA in Primary Cell':['',''],
        'CA in Secondary Cell':['',''],
        '5G: Intra-SgNB PSCell Change Success Rate' :['',''],
        'TH DL / Maximo  TH DL Diario 4G  > 7 Mbps ':['',''],
        'TH UL / TH UL 4G 0,5> Mbps':['',''],
        'TDD TH DL 5G':['',''],
        'TDD TH UL 5G':['',''],
        'FDD TH DL 5G':['',''],       
        'FDD TH UL 5G':['',''],                
        'Maximo TH DL 5G':['','Sin valores de TH DL en ','Teniendo en cuenta la cantidad de tráfico que cursan las celdas, se considera que el comportamiento es el esperado'],
        'Maximo TH UL 5G':['','Sin valores de TH en UL','Teniendo en cuenta la cantidad de tráfico que cursan las celdas, se considera que el comportamiento es el esperado'],
        '5G SgNB_Addition_Success_Rate':['',''],               
        '5G Average User Number':['',''],              
        '5G RLC DL Traffic (GB)':['',''],
        '5G RLC UL Traffic (GB)':['',''],
        'L.CSFB.E2W':['',''],
        'Maximo  TH DL Diario 4G ':['Valores entorno al target objetivo de TH','Valores bajos de TH en '],
        'TH UL 4G':['Valores entorno al target objetivo de TH (0.5 Mbps)','Valores bajos de TH en '],
        'NR Throughput DL User':['','','Teniendo en cuenta la cantidad de tráfico que cursan las celdas, se considera que el comportamiento es el esperado.'],
        'NR Throughput UL User':['','','Teniendo en cuenta la cantidad de tráfico que cursan las celdas, se considera que el comportamiento es el esperado.']}


def justification(cellList):
    text = ': '
    status = ''
    cellsNOK = []
    for i in cellList:
        
        if i[2] == "NOK":
            cellsNOK.append(i[0])
    
    if len(cellsNOK) > 0:
        cells = ', '.join(cellsNOK)
        status = 'NOK. ' + dic[i[1]][1] + cells + '. Ha quedado escalado al adjudicatario.'
    else:
        
        cont = 0
        for i in cellList:
            if i[2] != "NOK" and i[2] != "OK":
                cont += 1
                status = 'OK. ' + dic[i[1]][2] + '.'
            if cont == 0:
                status = 'OK. ' + dic[i[1]][0] + '.'    
                         
    text += status

    return text       

def table():
    pass

def create(listNOK,listNOKchecked,site,path,type):
    
    word = Document()
    
    now = time.strftime("%X")
    
    if now < str(12):
        word.add_paragraph(f"Buenos días,\nSe adjunta informe Babysitting 48 del site {site}. A continuación, se justifican sus KPI NOK.\n")
    else:
        word.add_paragraph(f"Buenas tardes,\nSe adjunta informe Babysitting 48 del site {site}. A continuación, se justifican sus KPI NOK.\n")

    if type == 'consolidation':
        table()

    for i in listNOK:
        total = ''
        cellList = []
        for j in listNOKchecked:
            if i == j[1]:
                cellList.append(j)      
        total += justification(cellList)
        p = word.add_paragraph()
        p.add_run(i).bold = True
        p.add_run(total)

        for k in range(len(cellList)):
            if i.find('>') == -1 and i.find('*') == -1:
                word.add_picture(f'.\\graphs\\{i}_{k+1}.png',width=Cm(20))
            elif i.find('>') != -1:
                element = i.replace('>','')
                word.add_picture(f'.\\graphs\\{element}_{k+1}.png',width=Cm(20))
            elif i.find('*') != -1:
                element = i.replace('*','')
                word.add_picture(f'.\\graphs\\{element}_{k+1}.png',width=Cm(20))
    name = path + '\\Babysitting_' + site + '.docx'
    word.save(name)
   
